import importlib.util
import json
import os
import tempfile
import unittest
from unittest.mock import patch

from langchain_core.documents import Document

from rag.bm25_retriever import BM25Retriever
from rag.document_parsers import (
    CsvParser,
    DocxParser,
    DocumentParserFactory,
    JsonParser,
    MarkdownParser,
    PdfParser,
    TxtParser,
    XlsxParser,
)
from rag.question_splitter import QuestionBasedSplitter
from utils.config_handler import chroma_conf
from utils.model_cache import ensure_docling_cache_dirs


HAS_DOCX = importlib.util.find_spec("docx") is not None
HAS_OPENPYXL = importlib.util.find_spec("openpyxl") is not None


class FakeLabel:
    def __init__(self, name: str):
        self.name = name


class FakeProv:
    def __init__(self, page_no: int):
        self.page_no = page_no


class FakeTextItem:
    def __init__(self, text: str, label_name: str, page_no: int):
        self.text = text
        self.label = FakeLabel(label_name)
        self.prov = [FakeProv(page_no)]


class FakeTableItem:
    def __init__(self, markdown: str, page_no: int):
        self.label = FakeLabel("TABLE")
        self.prov = [FakeProv(page_no)]
        self._markdown = markdown

    def export_to_markdown(self) -> str:
        return self._markdown


class FakeDoclingDocument:
    def __init__(self, items):
        self._items = items

    def iterate_items(self):
        for item, level in self._items:
            yield item, level


class FakeConversionResult:
    def __init__(self, items):
        self.document = FakeDoclingDocument(items)


class DocumentParserFactoryTests(unittest.TestCase):
    def test_factory_returns_expected_parsers(self):
        self.assertIsInstance(DocumentParserFactory.get_parser("sample.txt"), TxtParser)
        self.assertIsInstance(DocumentParserFactory.get_parser("sample.pdf"), PdfParser)
        self.assertIsInstance(DocumentParserFactory.get_parser("sample.csv"), CsvParser)
        self.assertIsInstance(DocumentParserFactory.get_parser("sample.md"), MarkdownParser)
        self.assertIsInstance(DocumentParserFactory.get_parser("sample.json"), JsonParser)
        self.assertIsInstance(DocumentParserFactory.get_parser("sample.docx"), DocxParser)
        self.assertIsInstance(DocumentParserFactory.get_parser("sample.xlsx"), XlsxParser)

    def test_factory_rejects_unsupported_extension(self):
        with self.assertRaises(ValueError):
            DocumentParserFactory.get_parser("sample.unsupported")


class ParserTests(unittest.TestCase):
    def test_pdf_parser_builds_docling_converter_with_ocr_auto_fallback(self):
        from docling.datamodel.base_models import InputFormat

        parser = PdfParser()
        converter = parser._build_converter()
        pdf_option = converter.format_to_options[InputFormat.PDF]
        pipeline_options = pdf_option.pipeline_options
        ocr_options = pipeline_options.ocr_options

        self.assertTrue(pipeline_options.do_ocr)
        self.assertTrue(pipeline_options.do_table_structure)
        self.assertEqual(type(ocr_options).__name__, "OcrAutoOptions")
        self.assertFalse(ocr_options.force_full_page_ocr)
        self.assertEqual(ocr_options.bitmap_area_threshold, 0.05)

    def test_pdf_parser_maps_docling_sections_pages_and_tables(self):
        parser = PdfParser()
        conversion_result = FakeConversionResult(
            [
                (FakeTextItem("Overview", "SECTION_HEADER", 1), 1),
                (FakeTextItem("Intro paragraph.", "PARAGRAPH", 1), 2),
                (FakeTextItem("Second page paragraph.", "PARAGRAPH", 2), 2),
                (FakeTextItem("Details", "SECTION_HEADER", 2), 2),
                (FakeTextItem("Detail body.", "PARAGRAPH", 2), 3),
                (FakeTableItem("| A | B |\n| - | - |\n| 1 | 2 |", 2), 3),
            ]
        )

        docs = parser._docling_result_to_documents("sample.pdf", conversion_result)

        self.assertEqual(len(docs), 4)
        self.assertEqual(docs[0].metadata["content_type"], "text")
        self.assertEqual(docs[0].metadata["page"], 1)
        self.assertEqual(docs[0].metadata["section_path"], "Overview")
        self.assertTrue(docs[0].metadata["ocr_enabled"])
        self.assertEqual(docs[0].metadata["ocr_mode"], "auto_fallback")
        self.assertEqual(docs[0].metadata["ocr_bitmap_area_threshold"], 0.05)
        self.assertIn("# Overview", docs[0].page_content)
        self.assertIn("Intro paragraph.", docs[0].page_content)

        self.assertEqual(docs[1].metadata["page"], 2)
        self.assertEqual(docs[1].metadata["section_title"], "Overview")
        self.assertIn("Second page paragraph.", docs[1].page_content)

        self.assertEqual(docs[2].metadata["section_path"], "Overview > Details")
        self.assertIn("## Details", docs[2].page_content)

        self.assertEqual(docs[3].metadata["content_type"], "table")
        self.assertEqual(docs[3].metadata["table_index"], 1)
        self.assertEqual(docs[3].metadata["page"], 2)
        self.assertEqual(docs[3].metadata["ocr_mode"], "auto_fallback")
        self.assertIn("| A | B |", docs[3].page_content)

    def test_csv_parser_creates_row_documents(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "records.csv")
            with open(file_path, "w", encoding="utf-8", newline="") as f:
                f.write("name,city\nAlice,Shanghai\n\nBob,Beijing\n")

            docs = CsvParser().parse(file_path)

        self.assertEqual(len(docs), 2)
        self.assertEqual(docs[0].metadata["file_type"], "csv")
        self.assertEqual(docs[0].metadata["row_index"], 2)
        self.assertIn("name: Alice", docs[0].page_content)
        self.assertIn("city: Beijing", docs[1].page_content)

    def test_json_parser_supports_list_dict_and_scalar(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            list_path = os.path.join(temp_dir, "records.json")
            dict_path = os.path.join(temp_dir, "settings.json")
            scalar_path = os.path.join(temp_dir, "scalar.json")

            with open(list_path, "w", encoding="utf-8") as f:
                json.dump([{"name": "Alice"}, {"name": "Bob"}], f, ensure_ascii=False)
            with open(dict_path, "w", encoding="utf-8") as f:
                json.dump({"profile": {"city": "Shanghai"}, "enabled": True}, f, ensure_ascii=False)
            with open(scalar_path, "w", encoding="utf-8") as f:
                json.dump("plain-value", f, ensure_ascii=False)

            list_docs = JsonParser().parse(list_path)
            dict_docs = JsonParser().parse(dict_path)
            scalar_docs = JsonParser().parse(scalar_path)

        self.assertEqual(len(list_docs), 2)
        self.assertEqual(list_docs[0].metadata["record_index"], 0)
        self.assertEqual(list_docs[1].metadata["json_path"], "$[1]")
        self.assertEqual(len(dict_docs), 2)
        self.assertEqual(dict_docs[0].metadata["json_path"], "$.profile")
        self.assertEqual(scalar_docs[0].metadata["json_path"], "$")
        self.assertEqual(scalar_docs[0].page_content, "plain-value")

    def test_markdown_parser_splits_by_headers(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "guide.md")
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("# Guide\nIntro\n\n## FAQ\nAnswer\n")

            docs = MarkdownParser().parse(file_path)

        self.assertEqual(len(docs), 2)
        self.assertEqual(docs[0].metadata["file_type"], "md")
        self.assertEqual(docs[1].metadata["markdown_path"], "Guide > FAQ")
        self.assertIn("Guide", docs[0].page_content)
        self.assertIn("FAQ", docs[1].page_content)

    @unittest.skipUnless(HAS_DOCX, "python-docx is not installed")
    def test_docx_parser_extracts_sections(self):
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "manual.docx")
            document = DocxDocument()
            document.add_heading("Overview", level=1)
            document.add_paragraph("System overview paragraph.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Field"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Version"
            table.cell(1, 1).text = "v1"
            document.save(file_path)

            docs = DocxParser().parse(file_path)

        self.assertEqual(len(docs), 1)
        self.assertEqual(docs[0].metadata["section_title"], "Overview")
        self.assertIn("System overview paragraph.", docs[0].page_content)
        self.assertIn("Field: Version", docs[0].page_content)
        self.assertIn("Value: v1", docs[0].page_content)

    @unittest.skipUnless(HAS_OPENPYXL, "openpyxl is not installed")
    def test_xlsx_parser_creates_row_documents(self):
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "sheet.xlsx")
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Customers"
            sheet.append(["name", "city"])
            sheet.append(["Alice", "Shanghai"])
            sheet.append([None, None])
            sheet.append(["Bob", "Beijing"])
            workbook.save(file_path)
            workbook.close()

            docs = XlsxParser().parse(file_path)

        self.assertEqual(len(docs), 2)
        self.assertEqual(docs[0].metadata["sheet_name"], "Customers")
        self.assertEqual(docs[0].metadata["row_index"], 2)
        self.assertIn("city: Beijing", docs[1].page_content)

    def test_ensure_docling_cache_dirs_uses_project_writable_defaults(self):
        with tempfile.TemporaryDirectory() as temp_dir, patch.dict(os.environ, {"DOC_MODEL_CACHE_DIR": temp_dir}, clear=False):
            cache_dirs = ensure_docling_cache_dirs()

            self.assertTrue(cache_dirs["base_dir"].startswith(temp_dir))
            self.assertTrue(os.path.isdir(cache_dirs["hub_cache"]))
            self.assertTrue(os.environ["DOC_MODEL_CACHE_DIR"].startswith(temp_dir))


class IntegrationShapeTests(unittest.TestCase):
    def test_question_splitter_preserves_source_metadata(self):
        splitter = QuestionBasedSplitter()
        documents = [
            Document(
                page_content="First structured paragraph.\nSecond structured paragraph.",
                metadata={
                    "source": "sample.json",
                    "file_type": "json",
                    "json_path": "$.profile",
                },
            )
        ]

        chunks = splitter.split_documents(documents)

        self.assertGreaterEqual(len(chunks), 1)
        self.assertEqual(chunks[0].metadata["file_type"], "json")
        self.assertEqual(chunks[0].metadata["json_path"], "$.profile")
        self.assertEqual(chunks[0].metadata["source"], "sample.json")

    def test_question_splitter_keeps_pdf_table_as_single_chunk(self):
        splitter = QuestionBasedSplitter()
        table_doc = Document(
            page_content="| A | B |\n| - | - |\n| 1 | 2 |",
            metadata={
                "source": "sample.pdf",
                "file_type": "pdf",
                "content_type": "table",
                "page": 1,
                "section_path": "Overview",
                "table_index": 1,
                "ocr_enabled": True,
                "ocr_mode": "auto_fallback",
                "ocr_bitmap_area_threshold": 0.05,
            },
        )

        chunks = splitter.split_documents([table_doc])

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].metadata["content_type"], "table")
        self.assertEqual(chunks[0].metadata["table_index"], 1)
        self.assertEqual(chunks[0].metadata["ocr_mode"], "auto_fallback")
        self.assertEqual(chunks[0].page_content, table_doc.page_content)

    def test_question_splitter_passthroughs_structured_pdf_text(self):
        splitter = QuestionBasedSplitter()
        text_doc = Document(
            page_content="# Overview\n\nIntro paragraph.\n\nMore content.",
            metadata={
                "source": "sample.pdf",
                "file_type": "pdf",
                "content_type": "text",
                "page": 2,
                "section_path": "Overview",
                "section_title": "Overview",
                "ocr_enabled": True,
                "ocr_mode": "auto_fallback",
                "ocr_bitmap_area_threshold": 0.05,
            },
        )

        chunks = splitter.split_documents([text_doc])

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].metadata["page"], 2)
        self.assertEqual(chunks[0].metadata["section_title"], "Overview")
        self.assertTrue(chunks[0].metadata["ocr_enabled"])
        self.assertEqual(chunks[0].page_content, text_doc.page_content)

    def test_bm25_retriever_uses_factory_for_supported_files(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            markdown_path = os.path.join(temp_dir, "faq.md")
            other_markdown_path = os.path.join(temp_dir, "shipping.md")
            with open(markdown_path, "w", encoding="utf-8") as f:
                f.write("# FAQ\nReturn policy details.\n")
            with open(other_markdown_path, "w", encoding="utf-8") as f:
                f.write("# Shipping\nDelivery details.\n")

            with patch.dict(
                chroma_conf,
                {"data_path": temp_dir, "allow_knowledge_file_type": ["md"]},
                clear=False,
            ):
                retriever = BM25Retriever()
                self.assertTrue(retriever.prepare_scope(blocking=True, timeout=5))
                retriever.retrieve("return policy")

        self.assertEqual(len(retriever.documents), 2)
        self.assertEqual(len(retriever.corpus), 2)
        self.assertIsNotNone(retriever.bm25)
        self.assertTrue(all(doc.metadata["file_type"] == "md" for doc in retriever.documents))


if __name__ == "__main__":
    unittest.main()
